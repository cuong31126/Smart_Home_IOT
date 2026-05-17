from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "Bao_cao_IoT_Smart_Home_MQTT_cap_nhat.docx"
OUTPUT = ROOT / "Bao_cao_IoT_Smart_Home_MQTT_cap_nhat_thiet_bi.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_table(table):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def replace_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            set_cell_text(table.cell(row_index, col_index), value, bold=row_index == 0)

    style_table(table)


def replace_single_cell_table(table, text):
    cell = table.cell(0, 0)
    cell.text = ""
    for index, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        if index == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
    shade_cell(cell, "F4F8FB")


def add_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_heading(doc, text):
    paragraph = doc.add_heading(text, level=1)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_note(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF2CC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(156, 87, 0)
    for line in lines:
        para = cell.add_paragraph(line)
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def main():
    doc = Document(INPUT)

    replacements = {
        "Cho phep nguoi dung chuyen Auto/Manual va dieu khien rieng den chinh GPIO2.":
            "Cho phep nguoi dung chuyen Auto/Manual va dieu khien 8 thiet bi trong nha qua dashboard.",
        "Dam bao LED canh bao luon phan anh dung trang thai cam bien, khong bi nguoi dung tat nham.":
            "Dam bao logic an toan: van gas tu dong khoa khi gas vuot nguong va khong cho mo thu cong khi dang canh bao.",
        "Kien truc moi dung y nghia dieu khien: nguoi dung chi dieu khien den chinh.":
            "Kien truc moi mo rong tu dieu khien den chinh sang nhom 8 thiet bi trong nha.",
        "Cac LED canh bao duoc bao ve khoi thao tac thu cong va luon phan anh cam bien.":
            "Cac thiet bi tu dong co the chay theo cam bien o Auto va duoc override o Manual khi an toan.",
        "Dashboard gon hon, tranh gay nham lan giua den canh bao va den nguoi dung dieu khien.":
            "Dashboard tach ro khu cam bien, khu thiet bi va canh bao GAS DANGER nhap nhay khi gas vuot nguong.",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]
        if text == "5. MQTT topics, Firebase paths va API":
            paragraph.paragraph_format.page_break_before = True

    table0 = doc.tables[0]
    for row in table0.rows:
        label = row.cells[0].text.strip()
        if label == "De tai":
            set_cell_text(row.cells[1], "He thong Smart Home doc cam bien va dieu khien 8 thiet bi tu dong")
        if label == "Ban cap nhat":
            set_cell_text(row.cells[1], "Them logic 8 thiet bi, Manual override, khoa van gas an toan va canh bao GAS DANGER nhap nhay")

    replace_single_cell_table(
        doc.tables[1],
        "Tom tat thay doi kien truc\nManual Mode khong chi dieu khien den chinh nua, ma cho phep dieu khien den, mai hien, cua so, quat, may hut am, coi nguoi la, coi gas va van gas. Rieng van gas bi khoa mo khi gas dang vuot nguong.",
    )

    replace_table(
        doc.tables[4],
        [
            ["Thiet bi", "GPIO", "Kieu du lieu", "Y nghia"],
            ["DHT22 SDA", "GPIO4", "Digital", "Nhiet do va do am"],
            ["MQ-2 AOUT", "GPIO34", "ADC", "Gia tri gas"],
            ["LDR AO", "GPIO35", "ADC", "Cuong do anh sang"],
            ["Rain button", "GPIO18", "Input pullup", "Mo phong trang thai mua"],
            ["PIR OUT", "GPIO19", "Digital", "Phat hien chuyen dong"],
            ["Den trong nha", "GPIO2", "Output", "Auto theo LDR hoac Manual tu dashboard"],
            ["Mai hien tu dong", "GPIO5", "Output", "Mo ra khi mua, co the override o Manual"],
            ["Coi canh bao nguoi la", "GPIO16", "Output", "Keu khi PIR phat hien chuyen dong"],
            ["Coi canh bao gas", "GPIO17", "Output", "Keu khi gas vuot nguong"],
            ["Quat lam mat", "GPIO21", "Output", "Bat khi nhiet do cao"],
            ["Cua so tu dong", "GPIO22", "Output", "Dong khi mua, mo khi an toan"],
            ["May hut am", "GPIO23", "Output", "Bat khi do am > 70%"],
            ["Van gas an toan", "GPIO25", "Output", "Khoa khi gas vuot nguong; khong cho mo Manual khi nguy hiem"],
        ],
    )

    replace_table(
        doc.tables[5],
        [
            ["Che do", "Nguon dieu khien", "Hanh vi"],
            ["Auto", "ESP32 xu ly theo cam bien", "Tu cap nhat 8 thiet bi: den, mai hien, cua so, quat, hut am, coi nguoi la, coi gas, van gas"],
            ["Manual", "Dashboard -> Backend REST -> MQTT -> ESP32", "Nguoi dung override tung thiet bi, nhung van gas khong duoc mo khi gasWarning = 1 hoac gas > 2000"],
        ],
    )

    replace_single_cell_table(
        doc.tables[6],
        "Nguyen tac an toan van gas\nKhi gas vuot nguong, he thong tu dong khoa van gas va bat canh bao. Dashboard, backend va ESP32 deu co lop chan rieng, nen lenh mo van gas se bi tu choi khi dang co GAS DANGER.",
    )

    replace_table(
        doc.tables[7],
        [
            ["Kenh", "Ten", "Huong", "Noi dung"],
            ["MQTT", "smarthome/sensors/data", "ESP32 -> Backend", "temperature, humidity, gas, light, rain, motion, trang thai 8 thiet bi, warning, mode"],
            ["MQTT", "smarthome/devices/status", "ESP32 -> Backend", "lamp, awning, window, fan, dehumidifier, securityAlarm, gasAlarm, gasValve, mode"],
            ["MQTT", "smarthome/mode/control", "Backend -> ESP32", "{ mode: auto | manual }"],
            ["MQTT", "smarthome/devices/control", "Backend -> ESP32", "Payload 0/1 cho tung thiet bi; vi du { awning: 1 }"],
            ["Firebase", "smarthome/current", "Backend -> Firebase -> Web", "Du lieu cam bien hien tai va trang thai output dang publish"],
            ["Firebase", "smarthome/devices/status", "Backend -> Firebase -> Web", "Trang thai 8 thiet bi va mode hien tai"],
            ["API", "POST /api/mode", "Web -> Backend", "Doi Auto/Manual"],
            ["API", "POST /api/control", "Web -> Backend", "Dieu khien thiet bi o Manual; chan mo van gas khi dang nguy hiem"],
        ],
    )

    replace_table(
        doc.tables[8],
        [
            ["File", "Noi dung da chinh"],
            ["frontend/src/lib/types.ts", "Mo rong DeviceControl va DeviceStatus cho 8 thiet bi"],
            ["frontend/src/lib/database.ts", "Normalize Firebase schema moi, fallback tu LED field cu neu du lieu cu con ton tai"],
            ["frontend/src/components/SmartHomeDashboard.tsx", "Panel Auto/Manual hien 8 thiet bi; GAS DANGER co animation nhap nhay"],
            ["backend/src/routes/controlRoutes.js", "POST /api/control chap nhan cac thiet bi moi va chan mo gasValve khi gas nguy hiem"],
            ["backend/src/mqttClient.js", "Backfill du lieu thiet bi khi MQTT payload con o format cu"],
            ["frontend/esp32/smart_home.ino", "Them logic Auto/Manual cho 8 output va khoa van gas an toan"],
            ["frontend/esp32/diagram.json", "Them LED mo phong cua so, may hut am va van gas tren Wokwi"],
        ],
    )

    replace_single_cell_table(
        doc.tables[11],
        "Demo 3: Dieu khien thiet bi trong nha tu dashboard\nNguoi dung bam Manual -> POST /api/mode -> ESP32 doi mode. Khi bam tung nut thiet bi -> POST /api/control -> backend publish smarthome/devices/control -> ESP32 cap nhat GPIO va publish lai smarthome/devices/status.",
    )

    replace_single_cell_table(
        doc.tables[12],
        "Demo 4: Logic tu dong va GAS DANGER\nO Auto, ESP32 tu bat/tat 8 thiet bi theo cam bien. Khi gas > 2000 hoac gasWarning = 1, dashboard hien GAS DANGER nhap nhay, coi gas bat va van gas bi khoa. Lenh mo van gas o Manual bi chan o ca UI, backend va ESP32.",
    )

    add_heading(doc, "9. Bo sung logic 8 thiet bi trong nha")
    p = doc.add_paragraph()
    p.add_run("He thong sau cap nhat khong con chi dieu khien den chinh. ").bold = True
    p.add_run(
        "Dashboard co mot panel Auto/Manual gom 8 thiet bi trong nha. O Auto, ESP32 tinh trang thai output tu cam bien. O Manual, nguoi dung co the tac dong tung thiet bi qua Firebase va MQTT, dong thoi van gas luon duoc uu tien an toan."
    )

    device_table = doc.add_table(rows=1, cols=4)
    headers = ["Thiet bi", "Field Firebase/MQTT", "Auto", "Manual"]
    for index, header in enumerate(headers):
        set_cell_text(device_table.cell(0, index), header, bold=True)

    rows = [
        ["Den trong nha", "lamp", "Bat khi light < 2000", "Bat/tat tu dashboard"],
        ["Mai hien tu dong", "awning", "Mo ra khi rain = 1", "Keo ra/thu vao"],
        ["Cua so tu dong", "window", "Dong khi mua, mo khi khong mua", "Mo/dong cua so"],
        ["Quat lam mat", "fan", "Bat khi temperature > 35 C", "Bat/tat quat"],
        ["May hut am", "dehumidifier", "Bat khi humidity > 70%", "Bat/tat may hut am"],
        ["Coi canh bao nguoi la", "securityAlarm", "Bat khi motion = 1", "Bat/tat coi"],
        ["Coi canh bao gas", "gasAlarm", "Bat khi gas > 2000", "Bat/tat coi gas"],
        ["Van gas an toan", "gasValve", "Khoa khi gas > 2000", "Chi duoc mo khi khong co GAS DANGER"],
    ]
    for row in rows:
        cells = device_table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    style_table(device_table)

    add_note(
        doc,
        "Luu y reset Firebase sau khi doi schema",
        [
            "Nen xoa nguyen node smarthome/devices/status va smarthome/control de tranh du lieu cu thieu field.",
            "Khong can xoa smarthome/logs neu muon giu lich su canh bao.",
            "Sau khi Wokwi va backend chay lai, status se duoc tao lai day du theo schema moi.",
        ],
    )

    for index, paragraph in list(enumerate(doc.paragraphs)):
        if index > 14 and "w:br" in paragraph._p.xml and not paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
